"""Tests for document management endpoints and pipeline.

Covers:
- Upload with valid PDF/DOCX
- Upload with invalid file type (rejected)
- Upload as non-admin (403)
- Document listing respects role-based filtering
- Document deletion cascades to chunks
- Document access level update re-stamps chunks
- Processing status endpoint
"""

import io
from unittest.mock import AsyncMock, patch, MagicMock
from uuid import uuid4

import pytest
import pytest_asyncio
from httpx import AsyncClient
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.document import Document
from app.models.chunk import DocumentChunk
from app.services.chunking_service import chunk_document, count_tokens
from app.utils.parsers import (
    detect_file_type,
    extract_text_from_pdf,
    extract_text_from_docx,
    ParsingError,
)
from tests.conftest import create_user, get_auth_headers

pytestmark = pytest.mark.asyncio


# ---------------------------------------------------------------------------
# Helpers: create fake PDF/DOCX bytes
# ---------------------------------------------------------------------------

def _make_minimal_pdf(text: str = "This is a test document with enough content for chunking.") -> bytes:
    """Create a minimal valid PDF with the given text using PyMuPDF."""
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def _make_minimal_docx(text: str = "This is a test DOCX document with enough content for chunking.") -> bytes:
    """Create a minimal valid DOCX with the given text using python-docx."""
    from docx import Document as DocxDocument
    doc = DocxDocument()
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Unit Tests: Parsers
# ---------------------------------------------------------------------------

class TestDetectFileType:
    """Tests for file type detection."""

    def test_detect_pdf(self):
        assert detect_file_type("report.pdf") == "pdf"

    def test_detect_pdf_uppercase(self):
        assert detect_file_type("REPORT.PDF") == "pdf"

    def test_detect_docx(self):
        assert detect_file_type("document.docx") == "docx"

    def test_detect_unsupported(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            detect_file_type("image.png")

    def test_detect_no_extension(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            detect_file_type("noextension")

    def test_detect_txt(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            detect_file_type("notes.txt")


class TestPDFParser:
    """Tests for PDF text extraction."""

    def test_extract_text_valid_pdf(self):
        pdf_bytes = _make_minimal_pdf("Hello World from PDF")
        text = extract_text_from_pdf(pdf_bytes)
        assert "Hello World from PDF" in text

    def test_extract_text_corrupted_pdf(self):
        with pytest.raises(ParsingError, match="Failed to open PDF"):
            extract_text_from_pdf(b"not a real pdf file")

    def test_extract_text_empty_pdf(self):
        """PDF with a blank page should raise ParsingError."""
        import fitz
        doc = fitz.open()
        doc.new_page()  # blank page, no text
        pdf_bytes = doc.tobytes()
        doc.close()

        with pytest.raises(ParsingError, match="no extractable text"):
            extract_text_from_pdf(pdf_bytes)


class TestDOCXParser:
    """Tests for DOCX text extraction."""

    def test_extract_text_valid_docx(self):
        docx_bytes = _make_minimal_docx("Hello World from DOCX")
        text = extract_text_from_docx(docx_bytes)
        assert "Hello World from DOCX" in text

    def test_extract_text_corrupted_docx(self):
        with pytest.raises(ParsingError, match="Failed to open DOCX"):
            extract_text_from_docx(b"not a real docx file")

    def test_extract_text_empty_docx(self):
        """DOCX with no paragraphs should raise ParsingError."""
        from docx import Document as DocxDocument
        doc = DocxDocument()
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

        with pytest.raises(ParsingError, match="no extractable text"):
            extract_text_from_docx(docx_bytes)


# ---------------------------------------------------------------------------
# Unit Tests: Chunking
# ---------------------------------------------------------------------------

class TestChunkingService:
    """Tests for document chunking."""

    def test_chunk_normal_text(self):
        """Normal text should produce multiple chunks."""
        text = "This is a sentence. " * 200  # Generate long text
        chunks = chunk_document(text)
        assert len(chunks) > 1
        for chunk in chunks:
            assert "content" in chunk
            assert "token_count" in chunk
            assert "chunk_index" in chunk
            assert chunk["token_count"] > 0

    def test_chunk_short_text(self):
        """Text shorter than chunk_size should return a single chunk."""
        text = "Short document."
        chunks = chunk_document(text)
        assert len(chunks) == 1
        assert chunks[0]["content"] == "Short document."
        assert chunks[0]["chunk_index"] == 0

    def test_chunk_empty_text(self):
        """Empty text should raise ValueError."""
        with pytest.raises(ValueError, match="empty"):
            chunk_document("")

    def test_chunk_whitespace_only(self):
        """Whitespace-only text should raise ValueError."""
        with pytest.raises(ValueError, match="empty"):
            chunk_document("   \n\n   ")

    def test_chunk_indexes_sequential(self):
        """Chunk indexes should be sequential starting from 0."""
        text = "This is a paragraph. " * 200
        chunks = chunk_document(text)
        indexes = [c["chunk_index"] for c in chunks]
        assert indexes == list(range(len(chunks)))

    def test_count_tokens(self):
        """Token counting should return positive integers for real text."""
        token_count = count_tokens("Hello, world! This is a test.")
        assert token_count > 0
        assert isinstance(token_count, int)


# ---------------------------------------------------------------------------
# Integration Tests: Document API Routes
# ---------------------------------------------------------------------------

class TestDocumentUpload:
    """Tests for POST /api/v1/documents/upload."""

    @patch("app.services.document_service.generate_embeddings_batch")
    async def test_upload_valid_pdf_as_admin(
        self,
        mock_embeddings: AsyncMock,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Admin can upload a valid PDF and it gets processed."""
        admin = await create_user(db_session, email="uploadadmin@test.com", role="admin")
        headers = get_auth_headers(admin)

        # Mock embedding generation (avoids OpenAI API call)
        mock_embeddings.return_value = [[0.1] * 3072]  # 1 chunk worth

        pdf_bytes = _make_minimal_pdf("This is a test document for upload.")

        response = await client.post(
            "/api/v1/documents/upload",
            headers=headers,
            files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
            data={"title": "Test Document", "access_level": "public"},
        )

        assert response.status_code == 201
        data = response.json()
        assert data["title"] == "Test Document"
        assert data["filename"] == "test.pdf"
        assert data["access_level"] == "public"
        assert data["processing_status"] == "completed"
        assert data["total_chunks"] >= 1

    @patch("app.services.document_service.generate_embeddings_batch")
    async def test_upload_valid_docx_as_admin(
        self,
        mock_embeddings: AsyncMock,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Admin can upload a valid DOCX."""
        admin = await create_user(db_session, email="uploaddocx@test.com", role="admin")
        headers = get_auth_headers(admin)

        mock_embeddings.return_value = [[0.1] * 3072]

        docx_bytes = _make_minimal_docx("Test DOCX content for upload testing.")

        response = await client.post(
            "/api/v1/documents/upload",
            headers=headers,
            files={"file": ("test.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"title": "Test DOCX", "access_level": "internal"},
        )

        assert response.status_code == 201
        data = response.json()
        assert data["title"] == "Test DOCX"
        assert data["access_level"] == "internal"
        assert data["processing_status"] == "completed"

    async def test_upload_invalid_file_type(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Uploading an unsupported file type returns 400."""
        admin = await create_user(db_session, email="uploadinvalid@test.com", role="admin")
        headers = get_auth_headers(admin)

        response = await client.post(
            "/api/v1/documents/upload",
            headers=headers,
            files={"file": ("image.png", b"fake image content", "image/png")},
            data={"title": "Invalid File", "access_level": "public"},
        )

        assert response.status_code == 400
        assert "unsupported" in response.json()["detail"].lower()

    async def test_upload_as_employee_forbidden(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Non-admin user gets 403 when trying to upload."""
        employee = await create_user(db_session, email="uploademployee@test.com", role="employee")
        headers = get_auth_headers(employee)

        pdf_bytes = _make_minimal_pdf("Test content")

        response = await client.post(
            "/api/v1/documents/upload",
            headers=headers,
            files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
            data={"title": "Forbidden Upload", "access_level": "public"},
        )

        assert response.status_code == 403

    async def test_upload_as_manager_forbidden(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Manager gets 403 when trying to upload."""
        manager = await create_user(db_session, email="uploadmanager@test.com", role="manager")
        headers = get_auth_headers(manager)

        pdf_bytes = _make_minimal_pdf("Test content")

        response = await client.post(
            "/api/v1/documents/upload",
            headers=headers,
            files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
            data={"title": "Forbidden Upload", "access_level": "public"},
        )

        assert response.status_code == 403


class TestDocumentListing:
    """Tests for GET /api/v1/documents — role-based filtering."""

    async def _seed_documents(
        self,
        db_session: AsyncSession,
        admin_id,
    ):
        """Create one document at each access level."""
        levels = ["public", "internal", "confidential", "restricted"]
        docs = []
        for level in levels:
            doc = Document(
                title=f"Doc {level.capitalize()}",
                filename=f"doc_{level}.pdf",
                file_size_bytes=1000,
                access_level=level,
                processing_status="completed",
                total_chunks=1,
                uploaded_by=admin_id,
            )
            db_session.add(doc)
            docs.append(doc)
        await db_session.commit()
        for doc in docs:
            await db_session.refresh(doc)
        return docs

    async def test_admin_sees_all_documents(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Admin sees documents at all 4 access levels."""
        admin = await create_user(db_session, email="listadmin@test.com", role="admin")
        await self._seed_documents(db_session, admin.id)
        headers = get_auth_headers(admin)

        response = await client.get("/api/v1/documents", headers=headers)
        assert response.status_code == 200
        data = response.json()
        levels = {d["access_level"] for d in data}
        assert "public" in levels
        assert "internal" in levels
        assert "confidential" in levels
        assert "restricted" in levels

    async def test_manager_sees_up_to_confidential(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Manager sees public, internal, confidential — NOT restricted."""
        admin = await create_user(db_session, email="listadmin2@test.com", role="admin")
        manager = await create_user(db_session, email="listmanager@test.com", role="manager")
        await self._seed_documents(db_session, admin.id)
        headers = get_auth_headers(manager)

        response = await client.get("/api/v1/documents", headers=headers)
        assert response.status_code == 200
        data = response.json()
        levels = {d["access_level"] for d in data}
        assert "restricted" not in levels
        assert "public" in levels
        assert "internal" in levels
        assert "confidential" in levels

    async def test_employee_sees_public_and_internal_only(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Employee sees only public and internal documents."""
        admin = await create_user(db_session, email="listadmin3@test.com", role="admin")
        employee = await create_user(db_session, email="listemployee@test.com", role="employee")
        await self._seed_documents(db_session, admin.id)
        headers = get_auth_headers(employee)

        response = await client.get("/api/v1/documents", headers=headers)
        assert response.status_code == 200
        data = response.json()
        levels = {d["access_level"] for d in data}
        assert "confidential" not in levels
        assert "restricted" not in levels
        assert "public" in levels
        assert "internal" in levels


class TestDocumentGetById:
    """Tests for GET /api/v1/documents/{id}."""

    async def test_get_document_accessible(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """User can access a document within their access level."""
        admin = await create_user(db_session, email="getadmin@test.com", role="admin")
        doc = Document(
            title="Accessible Doc",
            filename="accessible.pdf",
            file_size_bytes=1000,
            access_level="public",
            processing_status="completed",
            total_chunks=0,
            uploaded_by=admin.id,
        )
        db_session.add(doc)
        await db_session.commit()
        await db_session.refresh(doc)

        employee = await create_user(db_session, email="getemployee@test.com", role="employee")
        headers = get_auth_headers(employee)

        response = await client.get(f"/api/v1/documents/{doc.id}", headers=headers)
        assert response.status_code == 200
        assert response.json()["title"] == "Accessible Doc"

    async def test_get_document_access_denied(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Employee cannot access a confidential document (returns 404)."""
        admin = await create_user(db_session, email="getadmin2@test.com", role="admin")
        doc = Document(
            title="Secret Doc",
            filename="secret.pdf",
            file_size_bytes=1000,
            access_level="confidential",
            processing_status="completed",
            total_chunks=0,
            uploaded_by=admin.id,
        )
        db_session.add(doc)
        await db_session.commit()
        await db_session.refresh(doc)

        employee = await create_user(db_session, email="getemployee2@test.com", role="employee")
        headers = get_auth_headers(employee)

        response = await client.get(f"/api/v1/documents/{doc.id}", headers=headers)
        assert response.status_code == 404


class TestDocumentDeletion:
    """Tests for DELETE /api/v1/documents/{id}."""

    async def test_delete_cascades_to_chunks(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Deleting a document removes all its associated chunks."""
        admin = await create_user(db_session, email="deladmin@test.com", role="admin")
        headers = get_auth_headers(admin)

        # Create document with chunks
        doc = Document(
            title="To Delete",
            filename="delete_me.pdf",
            file_size_bytes=1000,
            access_level="public",
            processing_status="completed",
            total_chunks=2,
            uploaded_by=admin.id,
        )
        db_session.add(doc)
        await db_session.commit()
        await db_session.refresh(doc)

        # Create chunks for this document
        for i in range(2):
            chunk = DocumentChunk(
                document_id=doc.id,
                chunk_index=i,
                content=f"Chunk {i} content",
                token_count=5,
                access_level="public",
            )
            db_session.add(chunk)
        await db_session.commit()

        # Verify chunks exist
        result = await db_session.execute(
            select(func.count(DocumentChunk.id)).where(
                DocumentChunk.document_id == doc.id
            )
        )
        assert result.scalar() == 2

        # Delete the document
        response = await client.delete(
            f"/api/v1/documents/{doc.id}", headers=headers
        )
        assert response.status_code == 204

        # Verify document is gone
        result = await db_session.execute(
            select(Document).where(Document.id == doc.id)
        )
        assert result.scalar_one_or_none() is None

        # Verify chunks are cascaded away
        result = await db_session.execute(
            select(func.count(DocumentChunk.id)).where(
                DocumentChunk.document_id == doc.id
            )
        )
        assert result.scalar() == 0

    async def test_delete_nonexistent_returns_404(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Deleting a nonexistent document returns 404."""
        admin = await create_user(db_session, email="del404admin@test.com", role="admin")
        headers = get_auth_headers(admin)

        fake_id = uuid4()
        response = await client.delete(
            f"/api/v1/documents/{fake_id}", headers=headers
        )
        assert response.status_code == 404

    async def test_delete_as_employee_forbidden(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Non-admin cannot delete documents."""
        admin = await create_user(db_session, email="delemployeeadmin@test.com", role="admin")
        employee = await create_user(db_session, email="delemployee@test.com", role="employee")

        doc = Document(
            title="Protected",
            filename="protected.pdf",
            file_size_bytes=1000,
            access_level="public",
            processing_status="completed",
            total_chunks=0,
            uploaded_by=admin.id,
        )
        db_session.add(doc)
        await db_session.commit()
        await db_session.refresh(doc)

        headers = get_auth_headers(employee)
        response = await client.delete(
            f"/api/v1/documents/{doc.id}", headers=headers
        )
        assert response.status_code == 403


class TestDocumentUpdate:
    """Tests for PATCH /api/v1/documents/{id}."""

    async def test_update_access_level_restamps_chunks(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Changing a document's access_level updates all its chunks."""
        admin = await create_user(db_session, email="updateadmin@test.com", role="admin")
        headers = get_auth_headers(admin)

        doc = Document(
            title="Update Test",
            filename="update.pdf",
            file_size_bytes=1000,
            access_level="public",
            processing_status="completed",
            total_chunks=2,
            uploaded_by=admin.id,
        )
        db_session.add(doc)
        await db_session.commit()
        await db_session.refresh(doc)

        # Create chunks with the original access level
        for i in range(2):
            chunk = DocumentChunk(
                document_id=doc.id,
                chunk_index=i,
                content=f"Chunk {i}",
                token_count=3,
                access_level="public",
            )
            db_session.add(chunk)
        await db_session.commit()

        # Update access level to confidential
        response = await client.patch(
            f"/api/v1/documents/{doc.id}",
            headers=headers,
            json={"access_level": "confidential"},
        )
        assert response.status_code == 200
        assert response.json()["access_level"] == "confidential"

        # Verify chunks were re-stamped
        result = await db_session.execute(
            select(DocumentChunk).where(DocumentChunk.document_id == doc.id)
        )
        chunks = result.scalars().all()
        for chunk in chunks:
            assert chunk.access_level == "confidential"


class TestDocumentStatus:
    """Tests for GET /api/v1/documents/{id}/status."""

    async def test_get_status_completed(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Status endpoint returns processing status and chunk count."""
        admin = await create_user(db_session, email="statusadmin@test.com", role="admin")
        headers = get_auth_headers(admin)

        doc = Document(
            title="Status Test",
            filename="status.pdf",
            file_size_bytes=1000,
            access_level="public",
            processing_status="completed",
            total_chunks=5,
            uploaded_by=admin.id,
        )
        db_session.add(doc)
        await db_session.commit()
        await db_session.refresh(doc)

        response = await client.get(
            f"/api/v1/documents/{doc.id}/status", headers=headers
        )
        assert response.status_code == 200
        data = response.json()
        assert data["processing_status"] == "completed"
        assert data["total_chunks"] == 5

    async def test_get_status_nonexistent(
        self,
        client: AsyncClient,
        db_session: AsyncSession,
    ):
        """Status for nonexistent document returns 404."""
        admin = await create_user(db_session, email="status404admin@test.com", role="admin")
        headers = get_auth_headers(admin)

        fake_id = uuid4()
        response = await client.get(
            f"/api/v1/documents/{fake_id}/status", headers=headers
        )
        assert response.status_code == 404
