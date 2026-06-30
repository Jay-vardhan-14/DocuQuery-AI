"""Document file parsers for PDF and DOCX formats.

Uses PyMuPDF (fitz) for PDFs and python-docx for DOCX files.
Handles edge cases: empty documents, corrupted files, unsupported formats.
"""

import io
import logging
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Supported file extensions
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class ParsingError(Exception):
    """Raised when document parsing fails."""
    pass


def detect_file_type(filename: str) -> str:
    """Detect the file type from the filename extension.

    Args:
        filename: Original filename with extension.

    Returns:
        File type string: "pdf" or "docx".

    Raises:
        ValueError: If the file type is not supported.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    elif ext == ".docx":
        return "docx"
    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported formats: {', '.join(SUPPORTED_EXTENSIONS)}"
        )


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text content from a PDF file.

    Uses PyMuPDF (fitz) to extract text from all pages,
    preserving page order and basic structure.

    Args:
        file_bytes: Raw bytes of the PDF file.

    Returns:
        Extracted text content as a single string.

    Raises:
        ParsingError: If the PDF cannot be parsed or is empty.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        logger.error("Failed to open PDF: %s", str(e))
        raise ParsingError(f"Failed to open PDF file: {str(e)}")

    if doc.page_count == 0:
        doc.close()
        raise ParsingError("PDF file contains no pages")

    text_parts: list[str] = []

    for page_num in range(doc.page_count):
        try:
            page = doc[page_num]
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(page_text)
        except Exception as e:
            logger.warning(
                "Failed to extract text from PDF page %d: %s",
                page_num + 1,
                str(e),
            )
            # Continue processing remaining pages

    full_text = "\n\n".join(text_parts).strip()

    if not full_text:
        doc.close()
        raise ParsingError(
            "PDF file contains no extractable text content. "
            "The document may contain only images or scanned content."
        )

    logger.info(
        "Extracted %d characters from PDF (%d pages)",
        len(full_text),
        doc.page_count if hasattr(doc, 'page_count') else len(text_parts),
    )

    doc.close()

    return full_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text content from a DOCX file.

    Uses python-docx to extract all paragraph text,
    preserving paragraph structure.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        Extracted text content as a single string.

    Raises:
        ParsingError: If the DOCX cannot be parsed or is empty.
    """
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        logger.error("Failed to open DOCX: %s", str(e))
        raise ParsingError(f"Failed to open DOCX file: {str(e)}")

    text_parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text_parts: list[str] = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text_parts.append(cell_text)
            if row_text_parts:
                text_parts.append(" | ".join(row_text_parts))

    full_text = "\n\n".join(text_parts).strip()

    if not full_text:
        raise ParsingError(
            "DOCX file contains no extractable text content. "
            "The document may be empty or contain only images."
        )

    logger.info("Extracted %d characters from DOCX", len(full_text))

    return full_text


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from a document file based on its type.

    Convenience function that detects the file type and calls
    the appropriate parser.

    Args:
        file_bytes: Raw bytes of the document file.
        filename: Original filename with extension.

    Returns:
        Extracted text content.

    Raises:
        ValueError: If the file type is not supported.
        ParsingError: If parsing fails.
    """
    file_type = detect_file_type(filename)

    if file_type == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif file_type == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"No parser available for file type: {file_type}")
